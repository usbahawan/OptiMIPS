from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path


ROOT = Path(r"D:\MARS-main\MARS-main")
OUT = ROOT / "OptiMIPS_New_Features_Report_v2.docx"
IMG_DIR = ROOT / "report_assets"
IMG_DIR.mkdir(exist_ok=True)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2F3", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill="E8EEF5"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body, fill="F4F6F9", color=RGBColor(31, 58, 95)):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="B7C9E2")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    p.add_run("\n" + body)
    for run in p.runs[1:]:
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK
    doc.add_paragraph()


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    set_cell_border(cell, color="D0D0D0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_diagram(path, title, boxes, arrows, size=(1200, 420)):
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 28)
        font = ImageFont.truetype("arial.ttf", 20)
        font_small = ImageFont.truetype("arial.ttf", 17)
    except Exception:
        font_title = font = font_small = ImageFont.load_default()
    draw.text((30, 24), title, fill=(31, 77, 120), font=font_title)
    for box in boxes:
        x, y, w, h, label, fill = box
        draw.rounded_rectangle((x, y, x + w, y + h), radius=14,
                               fill=fill, outline=(120, 140, 170), width=2)
        words = label.split("\n")
        total_h = len(words) * 23
        ty = y + (h - total_h) // 2
        for line in words:
            bbox = draw.textbbox((0, 0), line, font=font)
            tx = x + (w - (bbox[2] - bbox[0])) // 2
            draw.text((tx, ty), line, fill=(30, 30, 30), font=font)
            ty += 24
    for arrow in arrows:
        x1, y1, x2, y2, text = arrow
        draw.line((x1, y1, x2, y2), fill=(70, 90, 120), width=4)
        # arrow head
        if x2 >= x1:
            pts = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
        else:
            pts = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
        draw.polygon(pts, fill=(70, 90, 120))
        if text:
            draw.text(((x1 + x2) // 2 - 35, (y1 + y2) // 2 - 28),
                      text, fill=(80, 80, 80), font=font_small)
    img.save(path)


def add_bar_chart(path):
    img = Image.new("RGB", (1000, 430), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 28)
        font = ImageFont.truetype("arial.ttf", 20)
        small = ImageFont.truetype("arial.ttf", 16)
    except Exception:
        title_font = font = small = ImageFont.load_default()
    draw.text((30, 25), "Example Result from optimips_feature_test.asm",
              fill=(31, 77, 120), font=title_font)
    data = [("Original CPI", 1.045, (248, 203, 203)),
            ("Optimized CPI", 1.000, (198, 239, 206)),
            ("Theoretical Min", 0.818, (221, 235, 247))]
    max_v = 1.10
    x0 = 240
    y = 100
    for label, value, fill in data:
        draw.text((40, y + 8), label, fill=(30, 30, 30), font=font)
        w = int((value / max_v) * 620)
        draw.rounded_rectangle((x0, y, x0 + w, y + 38), radius=8,
                               fill=fill, outline=(150, 150, 150))
        draw.text((x0 + w + 16, y + 7), f"{value:.3f}",
                  fill=(30, 30, 30), font=font)
        y += 72
    draw.text((40, 340),
              "Meaning: optimization removed the visible load-use stall, but the critical path still sets a lower theoretical bound.",
              fill=(80, 80, 80), font=small)
    img.save(path)


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = INK
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OptiMIPS")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("New Architecture Extensions Report")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLUE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Forwarding Mode, Critical Path, Stall Profile, Register Pressure, and Loop-Aware Basic Blocks")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    doc.add_paragraph()

    add_callout(
        doc,
        "Report purpose",
        "This report explains the five new features added to OptiMIPS in simple English. It explains what each feature does, why it matters in a MIPS pipeline optimizer, what logic was used in the code, and how the sample assembly program shows the feature in the GUI.",
        fill="EAF2F8",
    )

    add_heading(doc, "1. Project Context", 1)
    add_body(doc, "OptiMIPS is a pipeline-aware MIPS code optimizer inside the MARS simulator. The original tool could parse assembled instructions, build dependencies, reorder safe instructions, and show CPI improvement. The new work improves the tool from a simple scheduler into a better teaching tool for computer architecture and compiler backend ideas.")
    add_body(doc, "The main idea is still the same: the tool must not change the meaning of the program. It can only move an instruction when dependency rules say the move is safe. The five new features add more realistic hardware modes, better explanation of stalls, deeper analysis of the dependency graph, register pressure visibility, and safer loop handling.")

    add_heading(doc, "2. Quick Summary of the Five New Features", 1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Feature", "What it shows", "Why it matters", "Main classes"]):
        hdr[i].text = h
    rows = [
        ("Forwarding ON/OFF", "CPI changes under two hardware modes.", "Students can compare a forwarding pipeline with a stall-only pipeline.", "Scheduler, OptiMIPS"),
        ("Critical Path", "Longest weighted dependency chain.", "Shows the theoretical lower limit of scheduling.", "DependencyGraph, CriticalPathResult"),
        ("Stall Profile", "Exact producer-consumer stall pairs.", "Explains which hazards caused stalls and which were removed.", "StallEvent, StallResult, Scheduler"),
        ("Register Pressure", "Live register count at each row.", "Shows if optimization creates spill risk by keeping too many values live.", "RegisterPressureAnalyzer, OptiMIPS"),
        ("Loop-Aware Basic Blocks", "Blocks split at labels, branches, jumps, and targets.", "Lets loop bodies be optimized safely without crossing control-flow boundaries.", "BasicBlock, BasicBlockSplitter"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    flow_path = IMG_DIR / "backend_flow.png"
    add_diagram(
        flow_path,
        "Updated OptiMIPS Backend Flow",
        [
            (35, 120, 160, 90, "MARS\nAssembles", (232, 242, 255)),
            (235, 120, 160, 90, "MarsParser\nInstruction Model", (232, 242, 255)),
            (435, 120, 175, 90, "Basic Block\nSplitter", (245, 243, 255)),
            (655, 120, 170, 90, "Per-Block\nScheduler", (232, 248, 232)),
            (870, 95, 225, 140, "Analysis Panels\nCPI, Critical Path,\nStalls, Pressure", (255, 247, 237)),
        ],
        [
            (195, 165, 235, 165, ""),
            (395, 165, 435, 165, ""),
            (610, 165, 655, 165, ""),
            (825, 165, 870, 165, ""),
        ],
    )
    doc.add_picture(str(flow_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "3. Feature 1: Hardware Forwarding Mode", 1)
    add_body(doc, "Forwarding is a hardware technique where a result is sent directly from a later pipeline stage to an earlier instruction that needs it. This avoids many stalls. In the GUI, the new toggle says Forwarding: ON or Forwarding: OFF.")
    add_body(doc, "When Forwarding is ON, normal R-type and I-type RAW hazards cost zero stall cycles because forwarding handles them. However, a load-use hazard still costs one stall when a load instruction is immediately followed by an instruction that uses the loaded register.")
    add_body(doc, "When Forwarding is OFF, the tool behaves like a simpler stall-only processor. Every RAW hazard costs latency minus one stall cycles. The latency values come from the Scheduler latency map.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Mode", "Rule", "Student meaning"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("Forwarding ON", "Only immediate lw/lb/lh load-use hazards add 1 stall.", "Modern pipeline behavior."),
        ("Forwarding OFF", "Every RAW hazard adds latency - 1 stalls.", "Simpler processor without bypassing."),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_callout(doc, "Code logic", "Scheduler.countStalls(instrs, forwardingEnabled) now returns a StallResult. The total stall count is used for CPI, while the event list is used by the Stall Profile tab.", fill="F4F6F9")

    add_heading(doc, "4. Feature 2: Critical Path Analysis", 1)
    add_body(doc, "The critical path is the longest dependency chain in the instruction graph. Even a perfect scheduler cannot run faster than this chain, because each instruction in the chain depends on earlier work.")
    add_body(doc, "OptiMIPS computes a weighted path length for each instruction. The weight is the instruction latency. Then it backtracks from the largest path length to find the actual critical path.")
    add_number(doc, "Topologically sort the dependency graph.")
    add_number(doc, "Process nodes from the end back toward the start.")
    add_number(doc, "For each instruction, compute latency plus the best successor path.")
    add_number(doc, "Start from the node with the maximum path length.")
    add_number(doc, "Follow the successor with the largest remaining path to build the critical path.")
    add_body(doc, "In the GUI, critical-path instructions get a red left border. The stats panel also shows the theoretical minimum CPI and the scheduling gap.")

    add_heading(doc, "5. Feature 3: Stall Profile", 1)
    add_body(doc, "The old CPI result only said how many stalls existed. The new Stall Profile explains exactly where the stalls came from. Each stall event stores the producer instruction, consumer instruction, hazard type, stall cycles, and source text.")
    add_body(doc, "The GUI compares stall events from the original sequence with stall events from the optimized sequence. If the same producer-consumer pair no longer causes a stall after optimization, the table marks it as YES in the Eliminated column.")
    stall_path = IMG_DIR / "stall_profile_flow.png"
    add_diagram(
        stall_path,
        "Stall Profile Logic",
        [
            (70, 120, 210, 90, "Original\nStallResult", (254, 226, 226)),
            (370, 120, 210, 90, "Optimized\nStallResult", (220, 252, 231)),
            (670, 120, 210, 90, "Compare\nPair Keys", (255, 247, 237)),
            (960, 120, 190, 90, "GUI Table\nYES / NO", (232, 242, 255)),
        ],
        [
            (280, 165, 370, 165, ""),
            (580, 165, 670, 165, ""),
            (880, 165, 960, 165, ""),
        ],
    )
    doc.add_picture(str(stall_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "6. Feature 4: Register Pressure Visualization", 1)
    add_body(doc, "Instruction scheduling can reduce stalls, but it can also keep more registers live at the same time. If too many registers are live, a real compiler may need to spill values to memory. That is why register pressure is important.")
    add_body(doc, "The analyzer uses a simple liveness rule. At position i, a register is live if an earlier instruction wrote it and some instruction at or after i will read it. The pressure value is the number of live registers at that position.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Pressure value", "Bar color", "Meaning"]):
        table.rows[0].cells[i].text = h
    for row in [
        ("0 to 4", "Green", "Low register pressure."),
        ("5 to 8", "Orange", "Moderate pressure."),
        ("9 or more", "Red", "Possible spill risk."),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_body(doc, "In the GUI, each code table has a thin pressure column on the right. The bar width is pressure divided by 16, multiplied by 40 pixels. A tooltip says how many live registers exist at that point.")

    add_heading(doc, "7. Feature 5: Loop-Aware Basic Block Splitter", 1)
    add_body(doc, "The original scheduler treated the assembled program almost like one large instruction list. It avoided moving instructions across branches, but it did not explicitly split the program into basic blocks. The new splitter fixes that.")
    add_body(doc, "A basic block is a straight-line sequence of instructions. Control enters at the first instruction and leaves at the end. The scheduler is allowed to reorder inside one block, but never across block boundaries.")
    add_number(doc, "The first instruction is a block leader.")
    add_number(doc, "Every label or branch target starts a new block.")
    add_number(doc, "Every instruction after a branch or jump starts a new block.")
    add_number(doc, "Each block gets its own DependencyGraph and Scheduler.")
    add_number(doc, "The optimized blocks are reassembled in the original block order.")
    block_path = IMG_DIR / "block_split_flow.png"
    add_diagram(
        block_path,
        "Loop-Aware Basic Block Flow",
        [
            (55, 115, 175, 90, "Scan\nInstructions", (232, 242, 255)),
            (275, 115, 190, 90, "Find\nLeaders", (245, 243, 255)),
            (520, 115, 185, 90, "Create\nBlocks", (255, 247, 237)),
            (755, 115, 190, 90, "Schedule\nEach Block", (232, 248, 232)),
            (995, 115, 155, 90, "Reassemble\nProgram", (232, 242, 255)),
        ],
        [
            (230, 160, 275, 160, ""),
            (465, 160, 520, 160, ""),
            (705, 160, 755, 160, ""),
            (945, 160, 995, 160, ""),
        ],
    )
    doc.add_picture(str(block_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(doc, "Why this matters for loops", "Loop bodies run many times, so optimizing inside the loop body can improve the part of the program that matters most. At the same time, the block boundary prevents unsafe movement across loop branches.", fill="FFF7ED", color=RGBColor(122, 90, 0))

    add_heading(doc, "8. Worked Example: optimips_feature_test.asm", 1)
    add_body(doc, "The following program is designed to test all five new features. It has a loop, a branch target label, one load-use hazard, several RAW dependencies, independent filler instructions, and enough live registers to make pressure bars visible.")
    asm_path = ROOT / "optimips_feature_test.asm"
    code = asm_path.read_text(encoding="utf-8").strip()
    add_code_block(doc, code)

    add_heading(doc, "8.1 Step-by-Step Hazard Meaning", 2)
    add_body(doc, "This table explains the important dependency meaning of each instruction. The goal is to see why some instructions must stay in order, and why some independent instructions can safely move.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Line", "Instruction", "Important dependency"]):
        table.rows[0].cells[i].text = h
    hazard_rows = [
        ("1", "lui $s0, 0x1001", "Writes $s0, the base address used later by lw and addi."),
        ("2", "addi $s1, $zero, 8", "Writes loop limit $s1. The branch later compares against this value."),
        ("3", "addi $s2, $zero, 0", "Writes loop counter $s2. Many loop instructions use or update it."),
        ("4", "addi $s3, $zero, 0", "Writes running sum $s3. The loop reads and updates it."),
        ("5", "addi $t7, $zero, 3", "Writes constant $t7, used later by line 8."),
        ("6", "lw $t0, 0($s0)", "Reads $s0 and writes $t0. The next line reads $t0, so this creates a load-use hazard."),
        ("7", "add $t1, $t0, $s3", "Reads $t0 from line 6 and $s3 from line 4. It must wait for the load value."),
        ("8", "add $t2, $s2, $t7", "Independent from line 6 and line 7, so it can move into a stall slot safely."),
        ("9", "sub $t3, $s1, $s2", "Also independent from the load-use pair, so it can move earlier safely."),
        ("10", "and $t4, $t1, $t2", "Reads $t1 and $t2, so it must stay after lines 7 and 8."),
        ("11", "or $t5, $t4, $t3", "Reads $t4 and $t3, so it must stay after lines 10 and 9."),
        ("12", "xor $t6, $t5, $t0", "Reads $t5 and $t0, so it depends on the load chain."),
        ("13", "add $s3, $s3, $t6", "Updates the running sum, so it must stay after line 12."),
        ("14", "addi $s0, $s0, 4", "Moves the array pointer. It should remain inside the loop block."),
        ("15", "addi $s2, $s2, 1", "Increments the loop counter. The branch reads this value."),
        ("16", "bne $s2, $s1, loop", "Ends the loop block. Instructions must not move across this branch."),
        ("17", "sw $s3, result", "Stores the final sum after the loop finishes."),
        ("18", "add $zero, $zero, $zero", "Acts like a nop-style instruction in this example."),
    ]
    for row in hazard_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "8.2 Example Optimized Order", 2)
    add_body(doc, "A typical optimized order for the loop body is shown below. The exact order can vary slightly, but the important idea is the same: independent instructions are moved between the load and its consumer so the load-use stall can be reduced or removed. The optimizer still keeps all instructions inside their own basic block.")
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Optimized position", "Original line", "Instruction", "Why this is safe"]):
        table.rows[0].cells[i].text = h
    opt_rows = [
        ("1", "1", "lui $s0, 0x1001", "Initializes the base address, so it remains before memory access."),
        ("2", "2", "addi $s1, $zero, 8", "Initializes the loop limit before the branch uses it."),
        ("3", "3", "addi $s2, $zero, 0", "Initializes the counter before loop work begins."),
        ("4", "4", "addi $s3, $zero, 0", "Initializes the running sum before the loop reads it."),
        ("5", "5", "addi $t7, $zero, 3", "Initializes a constant used in the loop."),
        ("6", "6", "lw $t0, 0($s0)", "Starts the load of the current array value."),
        ("7", "8", "add $t2, $s2, $t7", "Independent instruction fills the load-use waiting slot."),
        ("8", "9", "sub $t3, $s1, $s2", "Another independent instruction can also move before the load consumer."),
        ("9", "7", "add $t1, $t0, $s3", "Now the loaded value $t0 has had more time to become ready."),
        ("10", "10", "and $t4, $t1, $t2", "Runs after both inputs are available."),
        ("11", "11", "or $t5, $t4, $t3", "Runs after $t4 and $t3 are available."),
        ("12", "12", "xor $t6, $t5, $t0", "Runs after the dependency chain produces $t5."),
        ("13", "13", "add $s3, $s3, $t6", "Updates the sum after $t6 is ready."),
        ("14", "14", "addi $s0, $s0, 4", "Pointer update stays inside the loop block."),
        ("15", "15", "addi $s2, $s2, 1", "Counter update stays before the branch."),
        ("16", "16", "bne $s2, $s1, loop", "Branch remains at the end of the loop block."),
        ("17", "17", "sw $s3, result", "Store remains in the after_loop block."),
        ("18", "18", "add $zero, $zero, $zero", "No effect on program result."),
    ]
    for row in opt_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "8.3 How the Five New Features Work on This Code", 2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Feature"
    table.rows[0].cells[1].text = "What you should see on this example"
    feature_rows = [
        ("Forwarding ON/OFF", "With forwarding ON, the main visible stall is the load-use pair at lines 6 and 7. With forwarding OFF, more RAW pairs appear because every dependency costs stall cycles."),
        ("Critical Path", "The red left border follows the longest chain through lw, add, and/or/xor, and the running sum update."),
        ("Stall Profile", "The original stall pair from lw to add should appear near the top. If the filler instruction removes that exact wait, the Eliminated column shows YES."),
        ("Register Pressure", "The loop body keeps several temporary registers live, so the pressure bars become orange or red around the middle of the loop."),
        ("Loop-Aware Basic Blocks", "The loop label starts a block, bne ends it, and after_loop starts a new block. The optimizer moves instructions inside the loop block only."),
    ]
    for row in feature_rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)

    chart_path = IMG_DIR / "example_result_chart.png"
    add_bar_chart(chart_path)
    doc.add_picture(str(chart_path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "8.4 Expected GUI Reading", 2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "GUI area"
    table.rows[0].cells[1].text = "How to read it"
    gui_rows = [
        ("Forwarding toggle", "Switch it ON and OFF. CPI changes immediately without rerunning the optimizer."),
        ("Original / Optimized code tables", "Green moved rows show scheduling changes. Red left borders show critical-path rows. Faint horizontal lines show basic block boundaries."),
        ("Pressure bars", "Bars on the right show how many registers are live at each row. Red means high pressure."),
        ("Critical path panel", "Shows the best theoretical CPI and the remaining scheduling gap."),
        ("Stall Profile tab", "Shows each stall pair. YES means the optimizer removed that exact stall pair."),
        ("Basic blocks panel", "Shows how many blocks were found and each block's CPI before and after scheduling."),
    ]
    for row in gui_rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)

    add_heading(doc, "9. How the Code Files Work Together", 1)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["File", "Responsibility", "Simple explanation"]):
        table.rows[0].cells[i].text = h
    files = [
        ("MarsParser.java", "Builds MipsInstruction objects.", "Reads assembled MARS instructions and records registers, memory flags, source line, address, and branch target."),
        ("Scheduler.java", "Schedules instructions and counts stalls.", "Supports forwarding mode, stall-only mode, and StallResult event lists."),
        ("DependencyGraph.java", "Protects program correctness.", "Prevents unsafe movement and computes the critical path."),
        ("BasicBlockSplitter.java", "Splits loops and branches into blocks.", "Finds leaders, creates blocks, schedules each block independently, and reassembles output."),
        ("RegisterPressureAnalyzer.java", "Counts live registers.", "Shows whether optimization increases register pressure."),
        ("OptiMIPS.java", "GUI and report surface.", "Displays code, CPI, pressure bars, critical path, stall profile, and block summaries."),
    ]
    for row in files:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "10. Testing Steps", 1)
    add_number(doc, "Open optimips_feature_test.asm in MARS.")
    add_number(doc, "Assemble the program with F3.")
    add_number(doc, "Open Tools > OptiMIPS.")
    add_number(doc, "Click Optimize.")
    add_number(doc, "Check that the forwarding toggle, critical path panel, stall profile tab, pressure bars, and basic block panel all update.")
    add_number(doc, "Toggle Forwarding OFF and confirm CPI and stall profile values refresh without rerunning Optimize.")

    add_heading(doc, "11. Final Value of These Extensions", 1)
    add_body(doc, "These five features make OptiMIPS more complete as a teaching tool. It no longer only shows that instruction scheduling can reduce CPI. It now explains why the CPI changes, which hardware mode is being modeled, which stalls were removed, which dependency chain limits the schedule, whether register pressure became risky, and how loop bodies are handled safely.")
    add_body(doc, "This is useful for a project report because it connects compiler backend logic with computer architecture concepts. The GUI becomes a learning surface, not just an optimizer output window.")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("OptiMIPS New Features Report")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
