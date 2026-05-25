from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from build_new_features_report import (
    ROOT,
    IMG_DIR,
    BLUE,
    DARK_BLUE,
    INK,
    MUTED,
    add_heading,
    add_body,
    add_bullet,
    add_number,
    add_callout,
    add_code_block,
    add_diagram,
    add_bar_chart,
    style_table,
)


OUT = ROOT / "OptiMIPS_Complete_Project_Report_Updated.docx"


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = INK
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
    return doc


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OptiMIPS")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Pipeline-Aware MIPS Code Optimizer for MARS")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Complete Updated Project Report")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    add_callout(
        doc,
        "What this report contains",
        "This report explains the full OptiMIPS project, including the original optimizer features and the five new architecture extensions. It uses simple English, diagrams, tables, and worked examples so the design is easy to understand and defend.",
        fill="EAF2F8",
    )

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Project area"
    table.rows[0].cells[1].text = "Included features"
    rows = [
        ("Original optimizer", "MARS parser, instruction model, dependency graph, list scheduler, CPI calculation, hazard log, explanation table, apply/export/report."),
        ("New extensions", "Forwarding mode, critical path, stall profile, register pressure, loop-aware basic block scheduling."),
        ("Testing material", "Simple RAW hazard example and full loop-aware feature test program."),
    ]
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)


def add_overview(doc):
    add_heading(doc, "1. Project Overview", 1)
    add_body(doc, "OptiMIPS is a tool added to the MARS MIPS simulator. Its purpose is to help students understand instruction scheduling, pipeline stalls, data hazards, and CPI improvement. The user writes and assembles MIPS code in MARS. Then OptiMIPS reads the assembled instructions and tries to reorder safe instructions to reduce pipeline waiting time.")
    add_body(doc, "This project is educational. It is not trying to become a full industrial compiler. Instead, it makes backend compiler ideas visible inside a familiar simulator. The GUI shows the original code, optimized code, hazards, explanations, stall profile, critical path, register pressure, and basic block structure.")
    add_callout(
        doc,
        "Main idea",
        "OptiMIPS only moves instructions when the dependency graph says it is safe. The optimizer must improve scheduling without changing the meaning of the MIPS program.",
        fill="F4F6F9",
    )


def add_backend_flow(doc):
    add_heading(doc, "2. Complete Backend Flow", 1)
    add_body(doc, "The backend starts after MARS has already assembled the program. This is useful because MARS has already resolved registers, instruction names, source lines, and branch targets.")

    path = IMG_DIR / "complete_backend_flow.png"
    add_diagram(
        path,
        "Complete OptiMIPS Backend Flow",
        [
            (30, 120, 145, 90, "MARS\nProgram", (232, 242, 255)),
            (205, 120, 150, 90, "Parser\nMipsInstruction", (232, 242, 255)),
            (385, 120, 155, 90, "Basic\nBlocks", (245, 243, 255)),
            (570, 120, 165, 90, "Dependency\nGraphs", (255, 247, 237)),
            (765, 120, 160, 90, "Scheduler", (232, 248, 232)),
            (955, 95, 200, 140, "GUI Analysis\nCPI, Hazards,\nPressure, Reports", (242, 244, 255)),
        ],
        [
            (175, 165, 205, 165, ""),
            (355, 165, 385, 165, ""),
            (540, 165, 570, 165, ""),
            (735, 165, 765, 165, ""),
            (925, 165, 955, 165, ""),
        ],
    )
    doc.add_picture(str(path), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_number(doc, "MARS assembles the program.")
    add_number(doc, "MarsParser converts each ProgramStatement into a MipsInstruction object.")
    add_number(doc, "BasicBlockSplitter separates the program at labels, branch targets, branches, and jumps.")
    add_number(doc, "Each block gets a DependencyGraph.")
    add_number(doc, "Scheduler reorders safe instructions inside each block only.")
    add_number(doc, "OptiMIPS updates the GUI with CPI, hazards, critical path, stall profile, register pressure, and block results.")


def add_original_features(doc):
    add_heading(doc, "3. Original Features", 1)
    add_heading(doc, "3.1 Instruction Model", 2)
    add_body(doc, "Every assembled instruction is represented as a MipsInstruction object. This object stores the instruction index, mnemonic, destination register, source registers, memory flags, branch flag, original source text, source line, and now also address information.")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Meaning"
    fields = [
        ("index", "Original assembled instruction number."),
        ("mnemonic", "Instruction name such as add, lw, bne, or sw."),
        ("dest", "Register written by the instruction, if any."),
        ("sources", "Registers read by the instruction."),
        ("isMemLoad / isMemStore", "Whether the instruction reads or writes memory."),
        ("isBranch", "Whether the instruction changes control flow."),
        ("originalText", "Source text shown in the GUI and reports."),
        ("sourceLine", "Original source line used by Apply to Editor."),
        ("address / branchTargetAddress", "Used by loop-aware basic block detection."),
    ]
    for row in fields:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)

    add_heading(doc, "3.2 Dependency Graph", 2)
    add_body(doc, "The dependency graph is the safety system. An edge from instruction A to instruction B means B must remain after A. Without this graph, the optimizer might move instructions in a way that changes the program result.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Dependency type", "Meaning", "Example"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("RAW", "Read after write. B reads a register written by A.", "add writes $t0, lw reads $t0."),
        ("WAR", "Write after read. B writes a register read by A.", "A reads $t1, B later writes $t1."),
        ("WAW", "Write after write. Both instructions write the same register.", "Two instructions write $s0."),
        ("MEM", "Memory operations are kept conservatively ordered.", "lw and sw are not freely swapped."),
        ("CTRL", "Branch and jump boundaries must be respected.", "Nothing unsafe moves across bne."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "3.3 List Scheduling", 2)
    add_body(doc, "The scheduler uses list scheduling. It repeatedly chooses from instructions whose required predecessors are already scheduled. The selected instruction is usually the one that adds the fewest stalls.")
    add_number(doc, "Start with an empty optimized list.")
    add_number(doc, "Find ready instructions whose dependency predecessors are already placed.")
    add_number(doc, "Estimate the stall cost of each ready instruction.")
    add_number(doc, "Pick the lowest-cost safe instruction.")
    add_number(doc, "Update register ready times and append the instruction.")
    add_number(doc, "Repeat until the block is complete.")

    add_heading(doc, "3.4 CPI and Original GUI Features", 2)
    add_body(doc, "CPI means cycles per instruction. Lower CPI is better. OptiMIPS calculates CPI as instruction count plus stalls, divided by instruction count.")
    add_code_block(doc, "CPI = (instruction_count + stall_count) / instruction_count\n\nImprovement = ((old_CPI - new_CPI) / old_CPI) * 100")
    add_bullet(doc, "Original Code table shows the assembled program before scheduling.")
    add_bullet(doc, "Optimized Code table shows the new instruction order.")
    add_bullet(doc, "Hazard Log explains dependencies found in the original program.")
    add_bullet(doc, "Easy Explanation gives simple-language reasons for moves.")
    add_bullet(doc, "Apply to Editor replaces the current editor order when safe.")
    add_bullet(doc, "Export Optimized writes an optimized assembly file.")
    add_bullet(doc, "Save Report writes a text report with CPI and hazard details.")


def add_new_features(doc):
    add_heading(doc, "4. New Features Added", 1)
    add_body(doc, "The updated OptiMIPS adds five new architecture features. These make the tool more realistic and easier to explain in a project demonstration.")

    add_heading(doc, "4.1 Forwarding ON/OFF", 2)
    add_body(doc, "Forwarding is a hardware feature that passes a result directly from one pipeline stage to another. With forwarding ON, most R-type and I-type RAW hazards do not need stalls. The important exception is an immediate load-use hazard, because a value loaded from memory is not ready soon enough for the very next instruction.")
    add_body(doc, "With forwarding OFF, the processor model is simpler. Every RAW hazard costs latency minus one stalls.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Hardware mode", "Stall rule", "GUI result"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("Forwarding ON", "Only immediate lw/lb/lh load-use hazards cost 1 stall.", "CPI usually improves."),
        ("Forwarding OFF", "Every RAW hazard costs latency - 1.", "CPI becomes higher and more stall events appear."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "4.2 Critical Path", 2)
    add_body(doc, "The critical path is the longest dependency chain. Even a perfect schedule cannot be faster than this chain, because each instruction in the chain waits for the previous one.")
    add_number(doc, "Topologically sort the dependency graph.")
    add_number(doc, "Compute pathLength[node] = latency[node] + best successor path.")
    add_number(doc, "Pick the node with the highest path length.")
    add_number(doc, "Backtrack through successors to list the full critical path.")
    add_body(doc, "In the GUI, critical-path rows have a red left border. The stats panel shows the theoretical minimum CPI and the gap above that minimum.")

    add_heading(doc, "4.3 Stall Profile", 2)
    add_body(doc, "The Stall Profile tab explains the exact stall pairs. A StallEvent records the producer instruction, consumer instruction, hazard type, stall cycles, and source text. This is more useful than only showing total stalls.")
    add_body(doc, "The optimizer compares original stall pairs against optimized stall pairs. If the same producer-consumer pair no longer stalls, the Eliminated column shows YES.")

    add_heading(doc, "4.4 Register Pressure", 2)
    add_body(doc, "Register pressure means how many registers are live at one point in the program. Scheduling can reduce stalls, but sometimes it keeps more values live at the same time. If pressure becomes too high, a real compiler may need to spill registers to memory.")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Pressure", "Color", "Meaning"]):
        table.rows[0].cells[i].text = h
    for row in [("0-4", "Green", "Low pressure"), ("5-8", "Orange", "Medium pressure"), ("9+", "Red", "Possible spill risk")]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_body(doc, "The thin bars on the right side of both code tables show pressure row by row. The tooltip tells the exact number of live registers.")

    add_heading(doc, "4.5 Loop-Aware Basic Blocks", 2)
    add_body(doc, "A basic block is a straight-line section of code. Instructions may be reordered inside a block, but they must not move across block boundaries. This is especially important for loops, because loop bodies run many times and contain branch targets.")
    add_number(doc, "First instruction is a leader.")
    add_number(doc, "Every label or branch target starts a block.")
    add_number(doc, "Every instruction after a branch or jump starts a block.")
    add_number(doc, "Each block is scheduled independently.")
    add_number(doc, "Optimized blocks are joined back in the same block order.")

    path = IMG_DIR / "new_features_flow.png"
    add_diagram(
        path,
        "New Feature Analysis Layer",
        [
            (50, 115, 180, 90, "Forwarding\nMode", (232, 242, 255)),
            (275, 115, 180, 90, "Critical\nPath", (255, 247, 237)),
            (500, 115, 180, 90, "Stall\nProfile", (254, 226, 226)),
            (725, 115, 180, 90, "Register\nPressure", (232, 248, 232)),
            (950, 115, 180, 90, "Basic\nBlocks", (245, 243, 255)),
        ],
        [
            (230, 160, 275, 160, ""),
            (455, 160, 500, 160, ""),
            (680, 160, 725, 160, ""),
            (905, 160, 950, 160, ""),
        ],
    )
    doc.add_picture(str(path), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_simple_example(doc):
    add_heading(doc, "5. Worked Example 1: Simple Scheduling Example", 1)
    add_body(doc, "This example is similar to the original project report. It shows RAW hazards and one independent instruction that can move upward to fill a stall.")
    code = """# OptiMIPS Simple Test Program
# This program has intentional RAW hazards.

.text
.globl main
main:
    add  $t0, $t1, $t2    # line 1: writes $t0
    lw   $t3, 0($t0)      # line 2: reads $t0
    sub  $t4, $t3, $t1    # line 3: reads $t3
    add  $t5, $t6, $t7    # line 4: independent, safe to move up
    mul  $t8, $t5, $t0    # line 5: reads $t5 and $t0
    sw   $t8, 4($t0)      # line 6: stores result
    add  $zero, $zero, $zero  # line 7: nop-style instruction"""
    add_code_block(doc, code)

    add_heading(doc, "5.1 Step-by-Step Hazard Meaning", 2)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Line", "Instruction", "Important dependency"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "add $t0, $t1, $t2", "Writes $t0. Later lw, mul, and sw need $t0."),
        ("2", "lw $t3, 0($t0)", "Reads $t0 from line 1 and writes $t3."),
        ("3", "sub $t4, $t3, $t1", "Reads $t3 from line 2, so this is a load-use RAW hazard."),
        ("4", "add $t5, $t6, $t7", "Independent from lines 1 to 3, so it can move earlier safely."),
        ("5", "mul $t8, $t5, $t0", "Reads $t5 and $t0, so it must wait for lines 4 and 1."),
        ("6", "sw $t8, 4($t0)", "Reads $t8 and $t0, so it must wait for lines 5 and 1."),
        ("7", "add $zero, $zero, $zero", "Acts like a nop in this example."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "5.2 Example Optimized Order", 2)
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Optimized position", "Original line", "Instruction", "Why this is safe"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "1", "add $t0, $t1, $t2", "It produces $t0, so it must stay early."),
        ("2", "4", "add $t5, $t6, $t7", "Independent instruction fills a stall slot."),
        ("3", "2", "lw $t3, 0($t0)", "Now $t0 has had more time to become ready."),
        ("4", "5", "mul $t8, $t5, $t0", "Inputs $t5 and $t0 are available."),
        ("5", "3", "sub $t4, $t3, $t1", "Moved after useful work, reducing the load-use wait."),
        ("6", "6", "sw $t8, 4($t0)", "Store remains after the values it needs."),
        ("7", "7", "nop-style add", "No effect on the result."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)


def add_full_feature_example(doc):
    add_heading(doc, "6. Worked Example 2: Full Feature Test Program", 1)
    add_body(doc, "The second example tests the new features together. It uses real/basic instructions instead of pseudo-instructions so MARS rows are cleaner.")
    code = (ROOT / "optimips_feature_test.asm").read_text(encoding="utf-8").strip()
    add_code_block(doc, code)

    add_heading(doc, "6.1 Step-by-Step Hazard Meaning", 2)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Line", "Instruction", "Important dependency"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "lui $s0, 0x1001", "Writes base address $s0 used by lw and pointer update."),
        ("2", "addi $s1, $zero, 8", "Writes loop limit $s1 used by bne."),
        ("3", "addi $s2, $zero, 0", "Writes loop counter $s2."),
        ("4", "addi $s3, $zero, 0", "Writes running sum $s3."),
        ("5", "addi $t7, $zero, 3", "Writes constant $t7 used by line 8."),
        ("6", "lw $t0, 0($s0)", "Loads array value into $t0."),
        ("7", "add $t1, $t0, $s3", "Reads $t0 immediately after lw, creating a load-use hazard."),
        ("8", "add $t2, $s2, $t7", "Independent filler candidate."),
        ("9", "sub $t3, $s1, $s2", "Independent filler candidate."),
        ("10", "and $t4, $t1, $t2", "Depends on lines 7 and 8."),
        ("11", "or $t5, $t4, $t3", "Depends on lines 10 and 9."),
        ("12", "xor $t6, $t5, $t0", "Depends on line 11 and the loaded value."),
        ("13", "add $s3, $s3, $t6", "Updates running sum after $t6 is ready."),
        ("14", "addi $s0, $s0, 4", "Moves to the next array element."),
        ("15", "addi $s2, $s2, 1", "Increments loop counter."),
        ("16", "bne $s2, $s1, loop", "Ends the loop block and jumps back if needed."),
        ("17", "sw $s3, result", "Stores final sum after loop."),
        ("18", "add $zero, $zero, $zero", "Nop-style instruction."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "6.2 Example Optimized Order", 2)
    add_body(doc, "A typical optimized loop order moves independent lines 8 and 9 between the load and its consumer. The branch still ends the loop block, so the optimizer does not move instructions across the loop boundary.")
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Optimized position", "Original line", "Instruction", "Why this is safe"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "1", "lui $s0, 0x1001", "Base address must be ready before memory access."),
        ("2", "2", "addi $s1, $zero, 8", "Loop limit must be ready before branch."),
        ("3", "3", "addi $s2, $zero, 0", "Counter must start before loop."),
        ("4", "4", "addi $s3, $zero, 0", "Sum must start before loop."),
        ("5", "5", "addi $t7, $zero, 3", "Constant must be ready before line 8."),
        ("6", "6", "lw $t0, 0($s0)", "Starts the load."),
        ("7", "8", "add $t2, $s2, $t7", "Independent work fills the load-use slot."),
        ("8", "9", "sub $t3, $s1, $s2", "Also independent from the load-use pair."),
        ("9", "7", "add $t1, $t0, $s3", "Now $t0 has had time to become ready."),
        ("10", "10", "and $t4, $t1, $t2", "Both inputs are ready."),
        ("11", "11", "or $t5, $t4, $t3", "Both inputs are ready."),
        ("12", "12", "xor $t6, $t5, $t0", "Depends on $t5 and $t0."),
        ("13", "13", "add $s3, $s3, $t6", "Updates sum after $t6."),
        ("14", "14", "addi $s0, $s0, 4", "Pointer update stays inside loop block."),
        ("15", "15", "addi $s2, $s2, 1", "Counter update stays before branch."),
        ("16", "16", "bne $s2, $s1, loop", "Branch remains at block end."),
        ("17", "17", "sw $s3, result", "After_loop block stores result."),
        ("18", "18", "add $zero, $zero, $zero", "No effect on program result."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "6.3 How the New Features Appear in the GUI", 2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Feature"
    table.rows[0].cells[1].text = "Expected result on this program"
    rows = [
        ("Forwarding mode", "Forwarding ON mostly shows the load-use stall. Forwarding OFF shows many more RAW stalls."),
        ("Critical path", "Red left border follows the longest dependency chain through the loop body."),
        ("Stall Profile", "Shows producer-consumer stall pairs and whether optimization eliminated them."),
        ("Register pressure", "Bars grow in the loop body where many temporary registers are live."),
        ("Basic blocks", "The loop and after_loop sections are separated by faint horizontal separators."),
    ]
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)

    chart_path = IMG_DIR / "complete_example_result_chart.png"
    add_bar_chart(chart_path)
    doc.add_picture(str(chart_path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_files_and_testing(doc):
    add_heading(doc, "7. Files and Classes Used", 1)
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["File", "Role", "Explanation"]):
        table.rows[0].cells[i].text = h
    rows = [
        ("OptiMIPS.java", "Main GUI", "Builds the interface and refreshes all analysis panels."),
        ("MarsParser.java", "Parser bridge", "Reads assembled MARS instructions and creates MipsInstruction objects."),
        ("MipsInstruction.java", "Instruction model", "Stores registers, flags, source text, source line, and branch address data."),
        ("DependencyGraph.java", "Safety and critical path", "Stores ordering edges and computes critical path."),
        ("Scheduler.java", "Scheduling and stall model", "Schedules ready instructions and returns StallResult."),
        ("StallEvent.java / StallResult.java", "Stall profile data", "Stores exact stall events and totals."),
        ("RegisterPressureAnalyzer.java", "Liveness analysis", "Counts live registers at each instruction position."),
        ("BasicBlock.java / BasicBlockSplitter.java", "Loop-aware scheduling", "Splits and schedules the program by basic block."),
        ("CriticalPathResult.java", "Critical path result", "Stores path lengths, critical path, min cycles, and min CPI."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "8. User Flow", 1)
    add_number(doc, "Open MARS using the updated Mars.jar.")
    add_number(doc, "Open or write a MIPS assembly program.")
    add_number(doc, "Assemble the program with F3.")
    add_number(doc, "Open Tools > OptiMIPS.")
    add_number(doc, "Click Optimize.")
    add_number(doc, "Read the original and optimized tables.")
    add_number(doc, "Use the Forwarding toggle to compare hardware modes.")
    add_number(doc, "Open Hazard Log, Easy Explanation, and Stall Profile tabs.")
    add_number(doc, "Review critical path, register pressure, and basic block summary.")
    add_number(doc, "Export, apply, or save a report if needed.")

    add_heading(doc, "9. Limitations", 1)
    add_bullet(doc, "The CPI model is educational and simplified. It is not a complete cycle-accurate CPU simulator.")
    add_bullet(doc, "Memory dependencies are conservative. This is safer for students but may miss some advanced optimizations.")
    add_bullet(doc, "Register pressure analysis is simple liveness, not full register allocation.")
    add_bullet(doc, "The optimizer schedules within basic blocks only. It does not perform global code motion across blocks.")
    add_bullet(doc, "Pseudo-instructions can expand into multiple instructions, so real/basic instructions are better for clean demonstrations.")

    add_heading(doc, "10. Final Conclusion", 1)
    add_body(doc, "OptiMIPS now demonstrates a complete educational backend pipeline. The original project showed instruction scheduling and CPI improvement. The updated project also explains hardware forwarding, critical path limits, exact stall causes, register pressure tradeoffs, and loop-safe basic block scheduling.")
    add_body(doc, "This makes the tool stronger for a project demonstration because each GUI result is connected to an architecture concept. A student can see not only that the code improved, but also why it improved and what tradeoffs remain.")


def build():
    doc = setup_doc()
    add_title_page(doc)
    add_overview(doc)
    add_backend_flow(doc)
    add_original_features(doc)
    add_new_features(doc)
    add_simple_example(doc)
    add_full_feature_example(doc)
    add_files_and_testing(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("OptiMIPS Complete Updated Project Report")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
